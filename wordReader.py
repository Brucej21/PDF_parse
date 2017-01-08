#word reader
# Import the module
#version 4 _ is to make the file name more automated and to tag changes with XXX.  
#This works send cleaning up now. 
from docx import Document
from docx.shared import Pt
from shutil import copyfile
import os
import csv
#import codecs
'''
data  = csv.reader(open("wordPatterns.csv"), delimiter=",")
fields = data.next()
'''
try:
              os.remove('OutDoc.docx')
              print("File Removed!")
except:
	print("No starting File to remove")

#http://stackoverflow.com/questions/123198/how-do-i-copy-a-file-in-python
#text2Bak.docx is the starting doc to look at, test2.docx is where the edits will land. 
copyfile('OrigDoc.docx', 'OutDoc.docx')

#document = Document('test2Bak.docx')
#document.save('test2.docx')

#https://newcircle.com/s/post/1572/python_for_beginners_reading_and_manipulating_csv_files
#source of CSV et al. http://web.uvic.ca/~gkblank/Blank's%20Writing%20Quirk%20List.pdf

#CSV holds the library of grammer patterns
f = open('wordPatternsv5.csv')
##f = codecs.open('wordPatterns4.csv', encoding='utf-8') # GOOD--gives you Unicode

csv_f = csv.reader(f, delimiter=',')

i =1

for row in csv_f:
	doc= Document('OutDoc.docx')
	#print i
	#i += 1
	for p in doc.paragraphs:
		if row[0] in p.text:
			print(row[0])
			#aaa = row[0]
			#print "aaa = " + aaa
			text = p.text.replace(row[0],'X '+row[0]+' X'+' Y ' +row[1]+' Y')
			p.text = text
			doc.save('OutDoc.docx')		
# doc.save(filename)
#doc.save('test7.docx')